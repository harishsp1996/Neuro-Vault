"""
Utility functions for HelperGPT
Text extraction, chunking, and AI response generation
"""
import os
import re
import json
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
from datetime import datetime
import PyPDF2
import docx
from openai import AsyncAzureOpenAI
from dotenv import load_dotenv

load_dotenv()
logger = logging.getLogger(__name__)

# Azure OpenAI configuration
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
GPT_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4")

# Initialize Azure OpenAI client
azure_openai_client = AsyncAzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))

async def extract_text_from_file(file_path: str) -> str:
    """Extract text content from various file formats"""
    try:
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.txt':
            return await extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return await extract_text_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return await extract_text_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return ""

async def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content.strip()
    except Exception as e:
        logger.error(f"Error reading TXT file: {str(e)}")
        return ""

async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                    continue
        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"Error reading PDF file: {str(e)}")
        return ""

async def extract_text_from_word(file_path: str) -> str:
    """Extract text from Word document"""
    try:
        doc = docx.Document(file_path)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"Error reading Word document: {str(e)}")
        return ""

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks for processing"""
    try:
        if not text or len(text.strip()) == 0:
            logger.warning("Empty text provided for chunking")
            return []
        
        text = clean_text(text)
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            if end < len(text):
                sentence_end = text.rfind('.', end - 200, end)
                if sentence_end != -1 and sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    para_break = text.rfind('\n\n', end - 200, end)
                    if para_break != -1 and para_break > start + chunk_size // 2:
                        end = para_break + 2
                    else:
                        word_end = text.rfind(' ', end - 100, end)
                        if word_end != -1 and word_end > start + chunk_size // 2:
                            end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        logger.info(f"Text chunked into {len(chunks)} pieces")
        return chunks
        
    except Exception as e:
        logger.error(f"Error chunking text: {str(e)}")
        return [text] if text else []

def clean_text(text: str) -> str:
    """Clean and normalize text content"""
    try:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    except Exception as e:
        logger.error(f"Error cleaning text: {str(e)}")
        return text

def determine_response_type(question: str, similar_docs: List[Dict[str, Any]]) -> str:
    """Determine response type based on question and document relevance"""
    try:
        # Lowercase the question for comparison
        question_lower = question.lower().strip()
        
        # Check for simple conversational keywords
        conversational_keywords = ['hello', 'hi', 'hey', 'thanks', 'thank you', 'sorry']
        if any(keyword in question_lower for keyword in conversational_keywords):
            return 'conversational'
        
        # If there are relevant documents, prioritize document-based response
        if similar_docs and similar_docs[0].get('similarity_score', 0) > 0.3:
            logger.info(f"Using document-based response (similarity: {similar_docs[0]['similarity_score']:.3f})")
            return 'document_based'
            
        # Default to document_based even without high similarity to force document-centric answers
        logger.info("Using document-based response (default)")
        return 'document_based'
            
    except Exception as e:
        logger.error(f"Error determining response type: {str(e)}")
        return 'document_based'

async def generate_response(question: str, similar_docs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Generate AI response using document context when available"""
    try:
        start_time = datetime.now()
        logger.info(f"Generating response for: '{question}' with {len(similar_docs)} similar docs")
        
        response_type = determine_response_type(question, similar_docs)
        
        # Prepare context from similar documents
        context_chunks = []
        source_info = []
        
        if similar_docs:
            for i, doc in enumerate(similar_docs[:3]):
                source_identifier = f"[{i+1}]"
                context_chunks.append(
                    f"{source_identifier} Document: {doc['filename']}\nContent: {doc['chunk_text']}"
                )
                source_info.append({
                    "document_id": doc["document_id"],
                    "filename": doc["filename"],
                    "team": doc["team"],
                    "project": doc["project"],
                    "relevance_score": doc["similarity_score"]
                })
        
        context = "\n\n---\n\n".join(context_chunks)
        
        if response_type == 'document_based' and context:
            system_prompt = """You are HelperGPT, an AI assistant for internal company documentation.
Your role is to help employees find information from uploaded company documents.

Instructions:
- Provide accurate answers **based only on the provided context**.
- Keep your answers concise and to the point.
- **Reference specific documents using the provided [number] tags** at the end of relevant sentences.
- If the answer is not in the provided context, state clearly and concisely that you don't have the information.
- **Do not use any external knowledge.** Your knowledge is strictly limited to the provided documents.
- Maintain a friendly and helpful tone.
- Do not repeat information, give lengthy descriptions, or provide a summary.
- The user can ask for more details if they need them."""
            
            user_prompt = f"""Based on the following company documentation, please answer this question concisely: {question}

Company document context:
{context}

Please provide a clear and brief answer based on this information."""
            
        else: # conversational responses for greetings
            system_prompt = """You are HelperGPT, a helpful and friendly AI assistant.
Your only function is to respond to simple greetings and pleasantries like 'hi', 'hello', 'thanks', and 'sorry'.
Keep your responses short and friendly.
For any other question, state clearly that you can only answer questions based on uploaded documents.
Do not provide any other information or engage in other conversations."""
            
            user_prompt = f"""User query: {question}

Please provide a brief, friendly response."""

        # Generate response using Azure OpenAI
        response = await azure_openai_client.chat.completions.create(
            model=GPT_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=256, # Reduced max_tokens for shorter answers
            temperature=0.7
        )

        answer = response.choices[0].message.content.strip()
        
        # *** MODIFICATION START ***
        # Check if the response type is document-based and if no context was found
        if response_type == 'document_based' and not context:
            answer = "I'm sorry, I don't have much information. Please contact your manager for more details."
        # *** MODIFICATION END ***
        
        response_time = (datetime.now() - start_time).total_seconds() * 1000
        confidence = calculate_confidence(similar_docs, response_type)

        show_sources = (
            response_type == 'document_based' and
            similar_docs and
            similar_docs[0].get('similarity_score', 0) > 0.2
        )

        result = {
            "answer": answer,
            "sources": source_info if show_sources else [],
            "confidence": confidence,
            "response_time_ms": int(response_time),
            "response_type": response_type
        }

        logger.info(f"Generated {response_type} response with {len(source_info)} sources")
        return result

    except Exception as e:
        logger.error(f"Error generating response: {str(e)}")
        return {
            "answer": "I'm sorry, I encountered an error while processing your question. Please try again or contact support.",
            "sources": [],
            "confidence": 0.0,
            "response_time_ms": 0,
            "response_type": "error"
        }

def calculate_confidence(similar_docs: List[Dict[str, Any]], response_type: str) -> float:
    """Calculate confidence score based on similarity scores and response type"""
    try:
        if response_type == 'document_based' and similar_docs:
            scores = [doc.get("similarity_score", 0.0) for doc in similar_docs[:3]]
            avg_score = sum(scores) / len(scores) if scores else 0.0
            confidence = min(avg_score * 1.2, 1.0)
            confidence = max(confidence, 0.3)
            return round(confidence, 2)
        elif response_type == 'conversational':
            return 0.7
        return 0.5
    except Exception as e:
        logger.error(f"Error calculating confidence: {str(e)}")
        return 0.5

async def validate_azure_openai_connection() -> bool:
    """Test Azure OpenAI connection"""
    try:
        response = await azure_openai_client.chat.completions.create(
            model=GPT_DEPLOYMENT,
            messages=[{"role": "user", "content": "Hello"}],
            max_tokens=10
        )
        return True
    except Exception as e:
        logger.error(f"Azure OpenAI connection test failed: {str(e)}")
        return False

def validate_question(question: str) -> bool:
    """Validate user question"""
    try:
        if not question or len(question.strip()) < 3:
            return False
        if len(question) > 1000:
            return False
        return True
    except Exception as e:
        logger.error(f"Error validating question: {str(e)}")
        return False

def format_file_size(size_bytes: int) -> str:
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.1f} {size_names[i]}"

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage"""
    try:
        filename = os.path.basename(filename)
        filename = re.sub(r'[<>:"/\|?*]', '_', filename)
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[:250] + ext
        return filename
    except Exception as e:
        return f"file_{datetime.now().strftime('%Y%m%d_%H%M%S')}"


